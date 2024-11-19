import os
import json
import streamlit as st
import openai
import random
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# .env 파일에서 OpenAI API 키 로드
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")

if not api_key:
    raise ValueError("API 키가 설정되지 않았습니다. .env 파일을 확인하세요.")

openai.api_key = api_key

@st.cache_data
def load_and_merge_json_files(directory):
    merged_data = []
    for filename in os.listdir(directory):
        if filename.endswith(".json"):  # JSON 파일만 처리
            file_path = os.path.join(directory, filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    if isinstance(data, list):
                        merged_data.extend(data)  # 리스트 데이터 병합
                    elif isinstance(data, dict):  # 단일 객체일 경우 리스트로 변환 후 병합
                        merged_data.append(data)
            except json.JSONDecodeError:
                st.error(f"'{filename}'은(는) 유효한 JSON 파일이 아닙니다.")
    return merged_data

def generate_event(goal, strategy, audience, budget, data, temperature=0.7, max_tokens=300):
    adjusted_max_tokens = random.randint(max(100, max_tokens - 30), min(800, max_tokens + 30))
    messages = [
        {"role": "system", "content": "You are an expert event planner specializing in cosmetics."},
        {"role": "user", "content": f"""
        I need a cosmetics-related event plan that achieves the following:

        - **Goal**: {goal}
        - **Strategy**: {strategy}
        - **Target Audience**: {audience}
        - **Budget**: {budget}

        Use the provided product data as a reference:
        {data}
        """}
    ]
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=adjusted_max_tokens,
        temperature=temperature
    )
    return response['choices'][0]['message']['content'].strip(), adjusted_max_tokens

def create_event_plan_doc_with_table(plan):
    """Word 문서에 표를 생성하고 내용을 채우는 함수."""
    doc = Document()
    doc.add_heading("이벤트 기획서", level=1)

    doc.add_paragraph("작성일자: 2023년 2월 13일")
    doc.add_paragraph("작성자: 마케팅 기획팀 김철수")

    # 표 생성
    table = doc.add_table(rows=6, cols=2, style='Table Grid')  # 6행 2열 테이블
    table.autofit = True

    # 표에 내용 입력
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"

    table.cell(1, 0).text = "목표"
    table.cell(1, 1).text = plan['goal']

    table.cell(2, 0).text = "대상"
    table.cell(2, 1).text = plan['audience']

    table.cell(3, 0).text = "전략"
    table.cell(3, 1).text = plan['strategy']

    table.cell(4, 0).text = "예산"
    table.cell(4, 1).text = plan['budget']

    table.cell(5, 0).text = "이벤트 기획안"
    table.cell(5, 1).text = plan['event_plan']

    # 스타일 및 정렬 설정
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 문서를 BytesIO로 저장하여 반환
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("Cosmetics Event Planner")
    st.write("Enter the details below to generate a customized event plan.")

    directory = "."
    cosmetics_data = load_and_merge_json_files(directory)

    if not cosmetics_data:
        st.error("No valid JSON files found in the directory.")
        return
    else:
        st.success(f"Loaded {len(cosmetics_data)} records from JSON files in the directory.")

    if "event_plan" not in st.session_state:
        st.session_state["event_plan"] = None

    with st.form("event_form"):
        goal = st.text_input("Goal (목표, 브랜드명 포함)", "")
        strategy = st.text_input("Strategy (전략)", "")
        audience = st.text_input("Target Audience (타겟층)", "")
        budget = st.text_input("Budget (예산)", "")
        temperature = st.slider("Creativity Level (Temperature)", 0.0, 1.0, 0.5)
        max_tokens = st.slider("Set Max Tokens", 100, 800, 300)
        submit = st.form_submit_button("Generate Event Plan")

    if submit:
        if not all([goal, strategy, audience, budget]):
            st.warning("All fields are required.")
            return

        brand_name = goal.split()[0] if goal else ""

        if not brand_name:
            st.warning("Please include a brand name in the goal.")
            return

        filtered_data = [item for item in cosmetics_data if brand_name.lower() in item.get('브랜드명', '').lower()]

        if not filtered_data:
            st.warning(f"No products found for the brand '{brand_name}'.")
            return

        data_str = "\n".join([f"{item.get('브랜드명', 'Unknown Brand')} - {item.get('제품명', 'Unknown Product')}" for item in filtered_data[:10]])

        event_plan, used_max_tokens = generate_event(goal, strategy, audience, budget, data_str, temperature, max_tokens)

        st.session_state["event_plan"] = {
            "goal": goal,
            "audience": audience,
            "strategy": strategy,
            "budget": budget,
            "event_plan": event_plan,
            "tokens_used": used_max_tokens,
        }

    if st.session_state["event_plan"]:
        plan = st.session_state["event_plan"]
        st.markdown("### 생성된 이벤트 계획:")
        st.markdown(f"1. **목표**:\n   {plan['goal']}")
        st.markdown(f"2. **대상**:\n   {plan['audience']}")
        st.markdown(f"3. **전략**:\n   {plan['strategy']}")
        st.markdown(f"4. **이벤트 기획안**:\n\n{plan['event_plan']}")
        st.markdown(f"5. **예산**:\n   {plan['budget']}")

        doc_buffer = create_event_plan_doc_with_table(plan)
        st.download_button(
            label="Download Event Plan as Word Document (With Table)",
            data=doc_buffer,
            file_name="event_plan_with_table.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
